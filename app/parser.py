"""
File Parser Module
Extracts clean text from PDF and DOCX files.
"""

import fitz  # PyMuPDF
from docx import Document
from pathlib import Path


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    text = []
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text.append(page.get_text("text"))
        doc.close()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    return "\n".join(text)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    text = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text.strip())
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
    return "\n".join(text)


def extract_text_from_file(file_path: str) -> str:
    """Auto-detect file type and extract text."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")